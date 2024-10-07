# app.py

from flask import Flask, render_template, request, redirect, url_for
import google.generativeai as genai
from docx import Document

app = Flask(__name__)

def get_gemini_response(input):
    model = genai.GenerativeModel('gemini-pro')
    response = model.generate_content(input)
    return response.text

# Function to configure Gemini AI model with the provided API key
def configure_gemini_api(api_key):
    genai.configure(api_key=api_key)

# Replace with your actual API key
API_KEY = 'AIzaSyANwNlGmVsaCSZKXSIHZZo0utCAwN6kZOY'
configure_gemini_api(API_KEY)

# Function to generate MCQs and save to a Word document
def generate_mcqs():
    mcq_list = []

    input_prompt = """
    Generate 25 different frequently asked aptitude questions in an interview. the response shud have a question, 
    4 options and one correct answer, shuffle the correct answer in unique way dont always put option A as correct 
    answer shuffle it(don't bold Question, Options, Correct Answer from below)

    Demo Response:
    Question (number): 
    Option A: 
    Option B: 
    Option C: 
    Option D: 
    Correct Answer: 
    """

    options = get_gemini_response(input_prompt)
    options = get_gemini_response(input_prompt)
    if options:
        mcq_list.append({
            "options": options
        })
    print(mcq_list)
    output_path = "generated_mcqs.docx"
    save_mcq_to_docx(mcq_list, output_path)

    return mcq_list, output_path

# Function to save MCQs to a Word document
def save_mcq_to_docx(mcq_list, output_path):
    doc = Document()

    for i, mcq in enumerate(mcq_list, start=1):
        doc.add_paragraph(mcq['options'])
        doc.add_paragraph("")  # Empty line for separation between questions

    doc.save(output_path)

# Home page
@app.route('/')
def index():
    return render_template('index.html')

# Generate MCQs and save to document
@app.route('/generate_mcqs', methods=['POST'])
def generate_mcqs_route():
    mcq_list, output_path = generate_mcqs()
    return redirect(url_for('show_mcqs'))

# Show generated MCQs
@app.route('/show_mcqs')
def show_mcqs():
    mcq_list, output_path = generate_mcqs()
    return render_template('show_mcqs.html', mcq_list=mcq_list)

# Evaluate MCQ answers and display results
@app.route('/evaluate_results', methods=['POST'])
def evaluate_results():
    correct_answers = request.form.getlist('correct_answers')
    user_answers = request.form.getlist('user_answers')

    correct_count = sum(1 for user_ans, correct_ans in zip(user_answers, correct_answers) if user_ans == correct_ans)
    total_questions = len(correct_answers)

    return render_template('results.html', correct_count=correct_count, total_questions=total_questions)

if __name__ == '__main__':
    app.run(debug=True)
