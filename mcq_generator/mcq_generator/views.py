import os
from django.shortcuts import render, redirect
from django.http import HttpResponse
import google.generativeai as genai
from docx import Document



def configure_gemini_api(api_key):
    genai.configure(api_key=api_key)


def get_gemini_response(input):
    model = genai.GenerativeModel('gemini-pro')
    response = model.generate_content(input)
    return response.text


def save_mcq_to_docx(mcq_list, output_path):
    doc = Document()
    for i, mcq in enumerate(mcq_list, start=1):
        doc.add_paragraph(mcq['options'])
        doc.add_paragraph("")  # Empty line for separation between questions
    doc.save(output_path)


def generate_mcqs(request):
    if request.method == 'POST':
        api_key = "AIzaSyDTenTn5lavhor_cot1odUWBJmScuR6G30"


        configure_gemini_api(api_key)

        input_prompt = """
            Generate 25 different frequently asked technical questions in an interview. the response shud have a question, 
            4 options and one correct answer, shuffle the correct answer in unique way

            Demo Response:
            Question (number): %
            Option A: %
            Option B: %
            Option C: %
            Option D: %
            Correct Answer: %
            (dont bold Question, Options, Correct Answer from above)
            """

        mcq_list = []
        options = get_gemini_response(input_prompt)
        if options:
            mcq_list.append({
                "options": options
            })

        output_path = os.path.join('mcq_generator', 'generated_mcqs.docx')
        save_mcq_to_docx(mcq_list, output_path)

        return render(request, 'result.html', {'output_path': output_path})


    return render(request, 'index.html')
