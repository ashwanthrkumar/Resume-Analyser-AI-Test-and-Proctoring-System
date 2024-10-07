import streamlit as st
import google.generativeai as genai
from docx import Document


# Configure Streamlit page settings
st.set_page_config(
    page_title="MCQ Generator",
    page_icon="📝",
    layout="centered",
)


# Function to configure Gemini AI model with the provided API key
def configure_gemini_api(api_key):
    genai.configure(api_key=api_key)


# Replace with your actual API key
API_KEY = 'AIzaSyANwNlGmVsaCSZKXSIHZZo0utCAwN6kZOY'
configure_gemini_api(API_KEY)


# Function to get response from Gemini AI
def get_gemini_response(input):
    model = genai.GenerativeModel('gemini-pro')
    response = model.generate_content(input)
    return response.text


# Function to read questions from uploaded Word file



# Function to save MCQs to a Word document
def save_mcq_to_docx(mcq_list, output_path):
    doc = Document()

    for i, mcq in enumerate(mcq_list, start=1):
        #doc.add_paragraph(f"{i}. {mcq['question']}")
        doc.add_paragraph(mcq['options'])
        doc.add_paragraph("")  # Empty line for separation between questions

    doc.save(output_path)

input_prompt = """
Generate 25 different frequently asked technical questions in an interview. the response shud have a question, 
4 options and one correct answer, shuffle the correct answer in unique way

Demo Response:
Question (number): %
Option A: %
Option B: %
Option C: %
Option D: %
Correct Answer: %
(dont bold Question, Options, Correct Answer from above)
"""

# Streamlit app
def main():
    st.title("MCQ Generator")

    if st.button("Generate MCQs") is not None:

        mcq_list = []

        options = get_gemini_response(input_prompt)
        if options:
            mcq_list.append({
                "options": options
            })

        output_path = "generated_mcqs.docx"
        save_mcq_to_docx(mcq_list, output_path)

        st.subheader("MCQs Generated and Saved!")
        st.markdown(f"Your MCQs have been saved to {output_path}. You can download it [here](./{output_path}).")

        st.markdown("---")
        st.header("Answer MCQs")
        st.markdown("Select the correct option for each question:")

        correct_answers = []
        user_answers = []

        for mcq in mcq_list:
            #st.subheader(mcq['question'])
            options = mcq['options'].strip().splitlines()
            selected_option = st.radio("Options", options)

            correct_answer = options[-1].split(': ')[1].strip()  # Extracting correct answer from options
            correct_answers.append(correct_answer)
            user_answers.append(selected_option)

            st.markdown("---")

        st.subheader("Evaluation Results")
        correct_count = sum(
            1 for user_ans, correct_ans in zip(user_answers, correct_answers) if user_ans == correct_ans)
        total_questions = len(correct_answers)
        st.markdown(f"You got {correct_count} out of {total_questions} questions correct!")


if __name__ == "__main__":
    main()
