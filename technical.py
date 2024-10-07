import streamlit as st
import google.generativeai as genai
from docx import Document
from dotenv import load_dotenv
import os
import json

# Load environment variables if using dotenv
load_dotenv()

# Configure Streamlit page settings
st.set_page_config(
    page_title="MCQ Generator",
    page_icon="📝",
    layout="centered",
)

# Function to configure Gemini AI model with the provided API key
def configure_gemini_api(api_key):
    genai.configure(api_key=api_key)

# Replace with your actual API key
API_KEY = 'AIzaSyDTenTn5lavhor_cot1odUWBJmScuR6G30'
configure_gemini_api(API_KEY)

# Function to get response from Gemini AI
def get_gemini_response(input):
    model = genai.GenerativeModel('gemini-pro')
    response = model.generate_content(input)
    return response.text

# Function to read questions from uploaded Word file
def read_questions_from_docx():
    doc_path = "questions.docx"  # Hardcoded path
    doc = Document(doc_path)
    questions = []

    for para in doc.paragraphs:
        text = para.text.strip()
        print(text)
        if text:
            questions.append(text)

    return questions

input_prompt = """
Generate 25 different frequently asked technical questions in an interview. the response shud have a question, 
4 options and a 
correct answer

Demo Response:
Question (number): %
Option A: %
Option B: %
Option C: %
Option D: %
Correct Answer: %
(dont bold Question, Options, Correct Answer from above)
"""

# Function to save MCQs to a Word document
def save_mcq_to_docx(mcq_list, output_path):
    doc = Document()

    for mcq in mcq_list:
        #doc.add_paragraph(f"Question: {mcq['question']}")
        options = mcq['options'].splitlines()
        for i, option in enumerate(options):
            doc.add_paragraph(f"{option.strip()}")

        # Assuming the correct answer is the last line in options
        # correct_answer = options[-1].strip() if options else ""
        # doc.add_paragraph(f"Correct Answer: {correct_answer}")
        # doc.add_paragraph("")  # Empty line for separation between questions

    doc.save(output_path)

# Streamlit app
st.title("MCQ Generator")
# uploaded_file = st.file_uploader("Upload Your Questions Document", type="docx", help="Please upload the DOCX file")

submit = st.button("Generate MCQs")
uploaded_file = Document('questions.docx')
if submit:
    if uploaded_file is not None:
        questions = read_questions_from_docx()
        mcq_list = []



        options = get_gemini_response(input_prompt)
        if options:
            mcq_list.append({
                "options": options
            })

        output_path = "generated_mcqs.docx"
        save_mcq_to_docx(mcq_list, output_path)

        st.subheader("MCQs Generated and Saved!")
        st.markdown(f"Your MCQs have been saved to {output_path}. You can download it [here](./{output_path}).")