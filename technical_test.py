from flask import Flask, render_template, request, redirect, url_for, make_response
import google.generativeai as genai
from docx import Document
import re
import cv2
import dlib
import numpy as np
from imutils import face_utils

app = Flask(__name__)
app.config['SECRET_KEY'] = 'your_secret_key'

# Global variables to store questions, options, and correct answers
questions = []
options = []
correct_answers = []

# Configure generative AI API
def configure_gemini_api(api_key):
    genai.configure(api_key=api_key)

# Replace with your actual API key
API_KEY = 'AIzaSyANwNlGmVsaCSZKXSIHZZo0utCAwN6kZOY'
configure_gemini_api(API_KEY)

# Function to get response from generative AI model
def get_gemini_response(input_prompt):
    model = genai.GenerativeModel('gemini-pro')
    response = model.generate_content(input_prompt)
    return response.text

# Function to generate MCQs using generative AI
def generate_mcqs():
    mcq_list = []
    input_prompt = """
    Generate 25 different frequently asked technical questions in an interview. the response should have a question, 
    4 options, and one correct answer. Shuffle the correct answer in a unique way; don't always put option A as the correct 
    answer, shuffle it (don't bold Question, Options, Correct Answer from below).

    Demo format response:
    Question (number): "What is a class"
    Option A: "a class is a object"
    Option B: "a class is a blueprint"
    Option C: "a class is a instance"
    Option D: "none"
    Correct Answer: B
    don't bold any text of the response
    """

    options = get_gemini_response(input_prompt)
    if options:
        mcq_list.append({
            "options": options
        })

    output_path = "generated_mcqs.docx"
    save_mcq_to_docx(mcq_list, output_path)

    return mcq_list, output_path

# Function to save MCQs to a Word document
def save_mcq_to_docx(mcq_list, output_path):
    doc = Document()
    for i, mcq in enumerate(mcq_list, start=1):
        doc.add_paragraph(mcq['options'])
        doc.add_paragraph("")  # Empty line for separation between questions
    doc.save(output_path)

# Function to read content from a Word document
def read_docx(file_path):
    doc = Document(file_path)
    paragraphs = [paragraph.text for paragraph in doc.paragraphs]
    return "\n".join(paragraphs)

# Function to parse MCQs from the generated document
def parsing():
    global questions, options, correct_answers
    file_path = 'generated_mcqs.docx'
    text = read_docx(file_path)

    questions = []
    options = []
    correct_answers = []

    pattern1 = re.compile(
        r'\*\*Question \d+:\*\*\n(.*?)\nOption A:\s(.*?)\nOption B:\s(.*?)\nOption C:\s(.*?)\nOption D:\s(.*?)\nCorrect Answer:\s(.*?)\n',
        re.DOTALL
    )

    # Define the alternative regex pattern
    pattern2 = re.compile(
        r'\d+\.\sQuestion:\s(.*?)\n\s*Option A:\s(.*?)\n\s*Option B:\s(.*?)\n\s*Option C:\s(.*?)\n\s*Option D:\s(.*?)\n\s*Correct Answer:\s(.*?)\n',
        re.DOTALL
    )

    pattern3 = re.compile(
        r'\d+\.\s(.*?)\n\s*Option A:\s(.*?)\n\s*Option B:\s(.*?)\n\s*Option C:\s(.*?)\n\s*Option D:\s(.*?)\n\s*Correct Answer:\s(.*?)\n',
        re.DOTALL
    )

    pattern4 = re.compile(
        r'\*\*Question\s+\d+:\*\*\s+(.*?)\s+\*\s*Option\s+A:\s(.*?)\s*\*\s*Option\s+B:\s(.*?)\s*\*\s*Option\s+C:\s(.*?)\s*\*\s*Option\s+D:\s(.*?)\s*\*\s*Correct\s+Answer:\s(.*?)\s*'
    )

    pattern5 = re.compile(
        r'\*\*Question\s+\d+:\*\*\s+(.*?)\s+Option\s+A:\s(.*?)\s+Option\s+B:\s(.*?)\s+Option\s+C:\s(.*?)\s+Option\s+D:\s(.*?)\s+\*\*Correct\s+Answer:\*\*\s+(.*?)\s+'
    )

    pattern6 = re.compile(
        r'\d+\.\s(.*?)\n\s*A:\s(.*?)\n\s*B:\s(.*?)\n\s*C:\s(.*?)\n\s*D:\s(.*?)\n\s*Correct Answer:\s(.*?)\n'
    )

    pattern7 = re.compile(
        r'Question \((\d+)\): "(.*?)"\s*Option [A-D]: "(.*?)"\s*Option [A-D]: "(.*?)"\s*Option [A-D]: "(.*?)"\s*Option [A-D]: "(.*?)"\s*Correct Answer: (.*?)\n'
    )

    pattern8 = re.compile(
        r'Question (\d+): "(.*?)"\s*Option [A-D]: "(.*?)"\s*Option [A-D]: "(.*?)"\s*Option [A-D]: "(.*?)"\s*Option [A-D]: "(.*?)"\s*Correct Answer: (.*?)\n'
    )

    pattern9 = re.compile(
        r'Question \(\d+\): "(.*?)"\s+Option [A-D]: "(.*?)"\s+Option [A-D]: "(.*?)"\s+Option [A-D]: "(.*?)"\s+Option [A-D]: "(.*?)"\s+Correct Answer: (.)\n'
    )

    # Find all matches using the first pattern
    matches = pattern1.findall(text)

    # If the first pattern doesn't match any, use the alternative pattern
    if not matches:
        matches = pattern2.findall(text)

    if not matches:
        matches = pattern3.findall(text)

    if not matches:
        matches = pattern4.findall(text)

    if not matches:
        matches = pattern5.findall(text)
    if not matches:
        matches = pattern6.findall(text)
    if not matches:
        matches = pattern7.findall(text)

    if not matches:
        matches = pattern8.findall(text)
    if not matches:
        matches = pattern9.findall(text)

    for match in matches:
        question = f"Question: {match[0]}"
        option_list = [match[1], match[2], match[3], match[4]]
        correct_answer = match[5]

        questions.append(question)
        options.append(option_list)
        correct_answers.append(correct_answer)

    # Map correct answers to their corresponding index
    correct_answers = [ord(answer) - ord('A') for answer in correct_answers]

    # Debugging: Print the number of questions parsed
    print(f"Total questions parsed: {len(questions)}")

    # Print details for further debugging
    print("Questions, options, and correct answers:")
    for i, question in enumerate(questions):
        print(f"Q{i+1}: {question}")
        for j, option in enumerate(options[i]):
            print(f"Option {chr(65+j)}: {option}")
        print(f"Correct Answer: {correct_answers[i]}\n")

# Initialize face detector and facial landmark predictor
detector = dlib.get_frontal_face_detector()
predictor = dlib.shape_predictor("shape_predictor_68_face_landmarks.dat")

# 3D model points for pose estimation
model_points = np.array([
    (0.0, 0.0, 0.0),             # Nose tip
    (0.0, -330.0, -65.0),        # Chin
    (-225.0, 170.0, -135.0),     # Left eye left corner
    (225.0, 170.0, -135.0),      # Right eye right corner
    (-150.0, -150.0, -125.0),    # Left mouth corner
    (150.0, -150.0, -125.0)      # Right mouth corner
], dtype="double")

# Camera parameters
size = (640, 480)  # Example size, change to your actual video feed size
focal_length = size[1]
center = (size[1] // 2, size[0] // 2)
camera_matrix = np.array([
    [focal_length, 0, center[0]],
    [0, focal_length, center[1]],
    [0, 0, 1]
], dtype="double")

# Distortion coefficients (assuming no lens distortion)
dist_coeffs = np.zeros((4, 1))

# Video capture initialization
cap = cv2.VideoCapture(0)

@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        return redirect(url_for('generate_mcqs_route'))
    return render_template('index.html')

@app.route('/generate_mcqs', methods=['GET'])
def generate_mcqs_route():
    return render_template('loading.html')

@app.route('/start_generating', methods=['GET'])
def start_generating():
    generate_mcqs()
    parsing()
    return redirect(url_for('take_test'))

@app.route('/take_test', methods=['GET', 'POST'])
def take_test():
    global cap, detector, predictor, model_points, camera_matrix, dist_coeffs

    if not questions:
        return redirect(url_for('generate_mcqs_route'))

    question_index = int(request.args.get('q', 0))

    if question_index >= len(questions):
        return redirect(url_for('finish'))

    options_dict = [{'index': i, 'option': option} for i, option in enumerate(options[question_index])]

    if request.method == 'POST':
        user_answer = int(request.form['option'])
        response = make_response(redirect(url_for('take_test', q=question_index + 1)))
        response.set_cookie(f'q-{question_index}', str(user_answer))
        return response

    # Function to monitor head pose and face detection
    def monitor_behavior():
        while True:
            ret, frame = cap.read()
            if not ret:
                break

            gray = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)
            rects = detector(gray, 0)

            for rect in rects:
                shape = predictor(gray, rect)
                shape = face_utils.shape_to_np(shape)

                image_points = np.array([
                    shape[30],     # Nose tip
                    shape[8],      # Chin
                    shape[36],     # Left eye left corner
                    shape[45],     # Right eye right corner
                    shape[48],     # Left mouth corner
                    shape[54]      # Right mouth corner
                ], dtype="double")

                success, rotation_vector, translation_vector = cv2.solvePnP(
                    model_points, image_points, camera_matrix, dist_coeffs)

                (nose_end_point2D, jacobian) = cv2.projectPoints(
                    np.array([(0.0, 0.0, 1000.0)]), rotation_vector, translation_vector, camera_matrix, dist_coeffs)

                p1 = (int(image_points[0][0]), int(image_points[0][1]))
                p2 = (int(nose_end_point2D[0][0][0]), int(nose_end_point2D[0][0][1]))

                cv2.line(frame, p1, p2, (0, 255, 255), 2)

                # Calculate Euler angles
                rmat, _ = cv2.Rodrigues(rotation_vector)
                angles, _, _, _, _, _ = cv2.RQDecomp3x3(rmat)

                angle_threshold = 15  # Degrees threshold to detect significant head movement
                danger_threshold = 40 # Degrees threshold to detect extreme head movement

                # Monitor yaw (rotation around the vertical axis)
                yaw_angle = angles[1]
                if abs(yaw_angle) > danger_threshold:
                    # Redirect to finish page or show alert
                    return redirect(url_for('finish'))

            cv2.imshow('Frame', frame)
            if cv2.waitKey(1) & 0xFF == ord('q'):
                break

        cap.release()
        cv2.destroyAllWindows()

    # Start monitoring behavior when the Flask app starts
    monitor_behavior()

    return render_template('test.html', question=questions[question_index], options=options_dict)

@app.route('/finish', methods=['GET'])
def finish():
    return render_template('finish.html')

if __name__ == '__main__':
    app.run(debug=True)
